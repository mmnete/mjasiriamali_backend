from docx import Document
from docx.shared import Inches
import os
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.text import WD_LINE_SPACING
from docx.shared import Mm

class DocumentEditorInterface:

    def __init__(self, document = None):
        self.doc = document

    def open_document(self, path: str):
        pass
    
    def add_image(self, img_path: str):
        pass

    def close_document(self):
        pass
    
    def get_page_width(self):
        pass

class WordDocumentEditor(DocumentEditorInterface):

    def open_document(self, path: str):
        self.doc = Document(path)
        print("Opened document:", path)
    
    def prepend_image(self, img_path: str):
        if self.doc is None:
            raise ValueError("Document not open")

        print(str(self.get_page_width()) + ' page width')
        paragraph = self.doc.paragraphs[0]
        paragraph.insert_paragraph_before().add_run().add_picture(img_path, width=self.get_page_width())
        # run = paragraph.add_run()
        # run.add_picture(img_path, width=self.get_page_width())

        # # Adjust paragraph spacing after the image
        # paragraph_format = paragraph.paragraph_format
        # paragraph_format.space_after = Pt(0)  # Set space after to 0 points
        # paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY  # Use exact line spacing

    def append_image(self, img_path: str):
        if self.doc is None:
            raise ValueError("Document not open")

        print(self.get_page_width() + ' page width')
        
        self.doc.add_picture(img_path, width=self.get_page_width())
        print("Added image:", img_path)

    def close_document(self, output_path: str):
        if self.doc is None:
            raise ValueError("Document not open")
        
        if os.path.exists(output_path):
            # If the file exists, delete it
            os.remove(output_path)
            print(f"Deleted existing file: {output_path}")
        
        self.doc.save(output_path)
        print("Closing file")
    
    def get_page_width(self):
        if self.doc is None or len(self.doc.sections) == 0:
            raise ValueError("Document not open or has no sections")

        # Get the width of the page in inches
        # section = self.doc.sections[0]
        # return int((section.page_width - section.left_margin - section.right_margin) / 36000)
        page_width_inches = self.doc.sections[0].page_width.inches - self.doc.sections[0].left_margin.inches - self.doc.sections[0].right_margin.inches 
        # Convert inches to EMUs (1 inch = 914400 EMUs)
        return int(page_width_inches * 914400)
        # Convert inches to EMUs (1 inch = 914400 EMUs)
        # return int(page_width_inches * 914400)

def add_banner_helper(doc):
    word_editor = WordDocumentEditor(Document(doc))
    word_editor.prepend_image("banner.png")
    word_editor.close_document("new_document.docx")
    return "new_document.docx"

